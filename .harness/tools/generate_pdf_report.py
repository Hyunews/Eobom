import os
import subprocess
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

# --- SINGLE AUTHORITATIVE REPORT DIRECTORY ---
# 2026-08-07 재구성: .harness/projects/eobom/reports → 저장소 루트 reports/
REPORTS_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), "../../reports"))
os.makedirs(REPORTS_DIR, exist_ok=True)

# --- 1. DOCX GENERATION FUNCTION ---
def create_docx_report(file_path):
    doc = docx.Document()
    
    # Page Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    # Color Palette
    NAVY = RGBColor(30, 41, 59)      # #1E293B
    GOLD = RGBColor(217, 119, 6)     # #D97706
    SLATE = RGBColor(15, 23, 42)     # #0F172A
    GRAY = RGBColor(100, 116, 139)   # #64748B
    
    # Set Normal Style Font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Malgun Gothic'
    font.size = Pt(10)
    font.color.rgb = RGBColor(51, 65, 85)
    
    # Helper: Set Cell Shading
    def set_cell_shading(cell, color_hex):
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)

    # Helper: Set Table Borders
    def set_table_borders(table, color="CBD5E1"):
        tblPr = table._tbl.tblPr
        borders = parse_xml(f'''
            <w:tblBorders {nsdecls("w")}>
                <w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>
                <w:bottom w:val="single" w:sz="6" w:space="0" w:color="{color}"/>
                <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>
                <w:insideV w:val="none"/>
                <w:left w:val="none"/>
                <w:right w:val="none"/>
            </w:tblBorders>
        ''')
        tblPr.append(borders)

    # --- COVER SECTION ---
    p_tag = doc.add_paragraph()
    r_tag = p_tag.add_run("전략 기획 및 시장 분석 보고서 (Confidential)")
    r_tag.font.size = Pt(9.5)
    r_tag.font.bold = True
    r_tag.font.color.rgb = GOLD
    
    p_title = doc.add_paragraph()
    r_title = p_title.add_run("국내 웰다잉·장례 테크 4대 서비스 현황 및 이어봄 플랫폼 경쟁력 분석 보고서")
    r_title.font.size = Pt(20)
    r_title.font.bold = True
    r_title.font.color.rgb = SLATE
    p_title.paragraph_format.space_after = Pt(4)
    
    p_sub = doc.add_paragraph()
    r_sub = p_sub.add_run("실제 4대 웹사이트(고이, 마음부고, iFA 엔딩플랜, 망고하다) 정밀 딥다이브 분석 및 이어봄 토탈 케어 차별화 전략")
    r_sub.font.size = Pt(11)
    r_sub.font.color.rgb = GRAY
    p_sub.paragraph_format.space_after = Pt(20)
    
    # Divider Line
    p_div = doc.add_paragraph()
    p_div_border = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:space="1" w:color="D97706"/></w:pBdr>')
    p_div._p.get_or_add_pPr().append(p_div_border)
    p_div.paragraph_format.space_after = Pt(20)

    # Summary Box
    p_box = doc.add_paragraph()
    p_box.paragraph_format.left_indent = Inches(0.2)
    p_box.paragraph_format.right_indent = Inches(0.2)
    p_box.paragraph_format.space_after = Pt(20)
    r_box_title = p_box.add_run("■ Executive Summary\n")
    r_box_title.font.bold = True
    r_box_title.font.color.rgb = GOLD
    r_box_body = p_box.add_run(
        "본 보고서는 대한민국 초고령사회 진입에 따른 웰다잉(Well-Dying) 및 사후 케어 시장의 실제 서비스 현황을 파악하기 위해, "
        "국내 대표 4대 서비스인 ①고이(goifuneral.co.kr), ②마음부고(mbugo.com), ③iFA 엔딩플랜(ifa.co.kr), ④망고하다(mangohada.com)의 "
        "웹사이트 및 서비스 기능을 정밀 조사하였습니다. 이를 기반으로 기존 파편화된 서비스의 한계를 극복하는 이어봄(LastLetter) 토탈 케어 플랫폼의 "
        "생전-응급-사후 3-Tier 이원화 보안 및 15개 올인원 웹 컴포넌트 경쟁 우위를 분석·검증합니다."
    )
    r_box_body.font.size = Pt(9.5)

    # --- SECTION 1 ---
    h1 = doc.add_heading("1. 국내 시장 환경 및 4대 영역 성장 배경", level=1)
    h1.runs[0].font.color.rgb = NAVY
    h1.runs[0].font.size = Pt(14)
    h1.runs[0].font.bold = True
    
    p = doc.add_paragraph("대한민국은 2025년을 기점으로 65세 이상 고령 인구 비율이 20%를 넘어서는 '초고령사회'에 진입하였습니다. 자녀 세대의 부양 부담 완화와 주체적인 삶의 마무리를 준비하는 웰다잉 문화가 확산되면서, 사후 케어 시장은 단순 오프라인 상조를 넘어 IT 테크 및 데이터 기반 서비스로 급격히 진화하고 있습니다.")
    
    doc.add_heading("1.1 기존 사후 시장의 3대 핵심 페인 포인트 (Pain Points)", level=2)
    p_bullet1 = doc.add_paragraph(style='List Bullet')
    p_bullet1.add_run("선불식 상조의 고액 예치금 및 폐업 위험: ").bold = True
    p_bullet1.add_run("기존 상조사는 수백만 원 상당의 예치금을 수년간 납입받아 유동성 리스크가 존재하며, 소비자 신뢰도 감소.")
    
    p_bullet2 = doc.add_paragraph(style='List Bullet')
    p_bullet2.add_run("장례 비용의 비대칭성 및 불투명성: ").bold = True
    p_bullet2.add_run("장례식장, 수의/관 용품, 봉안당/수목장 장지 비용의 명확한 사전 가격비교 체계 부재.")
    
    p_bullet3 = doc.add_paragraph(style='List Bullet')
    p_bullet3.add_run("생전 준비와 사후 케어의 단절: ").bold = True
    p_bullet3.add_run("엔딩노트 작성, 사전 연명의료 의향서, 상속세 시뮬레이션, 부고장 전송, 디지털 유품 정리가 개별 사이트에 파편화되어 있음.")

    # --- SECTION 2 ---
    h2 = doc.add_heading("2. 국내 4대 핵심 서비스 정밀 분석 (Live Website Deep Dive)", level=1)
    h2.runs[0].font.color.rgb = NAVY
    h2.runs[0].font.size = Pt(14)
    h2.runs[0].font.bold = True

    p = doc.add_paragraph("사용자가 제시한 국내 4개 주요 웹사이트의 실제 운영 서비스 구조, 핵심 모델 및 한계점을 분석한 상세 내역입니다.")

    # Competitor 1: Goi
    doc.add_heading("2.1 고이 (Goi) — goifuneral.co.kr", level=2)
    p = doc.add_paragraph()
    p.add_run("• 포지셔닝: ").bold = True
    p.add_run("투명한 가격 정찰제 기반의 디지털 장례 플랫폼 (고이장례연구소, 대표 이창수, 누적 120억 투자 유치, 2025년 매출 67억 원 달성).\n")
    p.add_run("• 핵심 경쟁력: ").bold = True
    p.add_run("'100원 상조' 후불제 모델(월 100원 납입으로 가입 시점 가격 평생 보장, 발생 후 정산), 미사용 물품 100% 반품/환불, 전국 장례식장 및 봉안당/수목장 실시간 맞춤 견적, 24시간 장례 컨시어지 및 발인/화장 동행 케어.\n")
    p.add_run("• 한계점: ").bold = True
    p.add_run("사후 장례 의전 및 장지 시설 견적 비교에 치중되어 있어, 생전 응급 QR 카드, 256-bit 자산 암호화 금고, 상속세 자동 시뮬레이션, 디지털 유품/SNS 계정 정리 기능은 미제공.")

    # Competitor 2: mbugo
    doc.add_heading("2.2 마음부고 (mbugo) — mbugo.com", level=2)
    p = doc.add_paragraph()
    p.add_run("• 포지셔닝: ").bold = True
    p.add_run("무료 모바일 부고장 & 부고 문자 카카오톡 알림 서비스 (운영사 무문, 대표 문명진).\n")
    p.add_run("• 핵심 경쟁력: ").bold = True
    p.add_run("회원가입 및 앱 설치 없이 1분 만에 부고장/감사장 무료 생성, 카카오톡 공유 및 문자 URL 전송, 내 주변 장례식장 직통 전화 연결 및 지도 안내.\n")
    p.add_run("• 한계점: ").bold = True
    p.add_run("단순 '모바일 부고장 작성 및 공유' 도구에 한정되며, 엔딩노트, 장지 선택, 상속 세무, 사후 자산/계정 정리는 전무함.")

    # Competitor 3: iFA Ending Plan
    doc.add_heading("2.3 iFA 엔딩플랜 — ifa.co.kr", level=2)
    p = doc.add_paragraph()
    p.add_run("• 포지셔닝: ").bold = True
    p.add_run("대형 인슈테크 GA 기반 상속·증여 전문 컨설팅 및 재무설계 (아이에프에이 주식회사, 전국 2,200명 조직, 특허 4건 보유).\n")
    p.add_run("• 핵심 경쟁력: ").bold = True
    p.add_run("상속·증여 전문 연구소 운영, 서울 평균 아파트 실거래가(15.5억 원) 기준 상속세 모의 계산기 (Tax_Main.aspx), 상속재산 평가부터 절세/재원 마련(종신보험 레버리지) 및 유언장 공증까지 원스톱 마무리.\n")
    p.add_run("• 한계점: ").bold = True
    p.add_run("보험 상품 판매 및 오프라인 상속·증여 영업 마케팅 도구 성격이 강하며, 디지털 유산 정리, 사전 연명의료 응급 QR 카드, 장지 비교 기능 부재.")

    # Competitor 4: Mangohada
    doc.add_heading("2.4 망고하다 — mangohada.com", level=2)
    p = doc.add_paragraph()
    p.add_run("• 포지셔닝: ").bold = True
    p.add_run("유언과 삶의 회고를 담는 스토리텔링 기반 웰다잉 전문 스타트업 플랫폼.\n")
    p.add_run("• 핵심 경쟁력: ").bold = True
    p.add_run("감성적 유언 기록, 버킷리스트 작성, 장례 의전 전문 기업(엔딩스케치)과 MOU를 통한 온라인 기록 ➔ 실물 장례 실행 연계.\n")
    p.add_run("• 한계점: ").bold = True
    p.add_run("정서적 메시지 기록 위주로 운영되며, 256-bit AES 암호화 금고, 2인 유족 승인(Multi-Sig) 기술, 상속세 계산기 및 장지 맞춤 비교 등 기술적·기능적 시스템 미흡.")

    # --- SECTION 3 ---
    h3 = doc.add_heading("3. 서비스 기능 및 경쟁력 종합 비교 (Feature Comparison Matrix)", level=1)
    h3.runs[0].font.color.rgb = NAVY
    h3.runs[0].font.size = Pt(14)
    h3.runs[0].font.bold = True

    # Table Creation
    table = doc.add_table(rows=8, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    
    headers = ["비교 핵심 기능", "이어봄 (본 서비스)", "고이 (Goi)", "마음부고 (mbugo)", "iFA 엔딩플랜", "망고하다"]
    for i, title in enumerate(headers):
        cell = table.cell(0, i)
        cell.paragraphs[0].text = title
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, "1E293B")
        
    matrix_data = [
        ["서비스 통합 범위", "생전~사후 올인원", "사후 의전/장지 중심", "모바일 부고장 전용", "상속세/보험 중심", "생전 유언/의전 연계"],
        ["생전/응급 의료 QR 카드", "지원 (사전연명의료)", "미지원", "미지원", "미지원", "미지원"],
        ["256-bit AES & Multi-Sig", "지원 (유족 2인 승인)", "미지원", "미지원", "미지원 (일반 저장)", "미지원 (일반 저장)"],
        ["전국 장지/봉안당 비교", "지원 (전국 8대 권역)", "지원 (실시간 견적)", "미지원 (식장 검색만)", "미지원", "미지원"],
        ["상속세 자동 시뮬레이터", "지원 (자산/부채 차트)", "미지원", "미지원", "지원 (모의 계산기)", "미지원"],
        ["디지털 유품 (SNS) 정리", "지원 (3/8건 트래킹)", "미지원", "미지원", "미지원", "미지원"],
        ["1:1 전문가 상담 챗", "지원 (실시간 예약)", "24h 콜센터", "미지원", "오프라인 컨설팅", "미지원"]
    ]
    
    for row_idx, row_data in enumerate(matrix_data, start=1):
        for col_idx, cell_value in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.paragraphs[0].text = cell_value
            if col_idx == 1: # 이어봄 column highlight
                set_cell_shading(cell, "EFF6FF")
                cell.paragraphs[0].runs[0].font.bold = True
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(3, 105, 161)

    doc.add_paragraph().paragraph_format.space_after = Pt(15)

    # --- SECTION 4 ---
    h4 = doc.add_heading("4. 이어봄 플랫폼의 4대 독보적 차별화 경쟁 우위", level=1)
    h4.runs[0].font.color.rgb = NAVY
    h4.runs[0].font.size = Pt(14)
    h4.runs[0].font.bold = True

    p = doc.add_paragraph()
    p.add_run("1. 생전-응급-사후 3-Tier 이원화 보안 체계 (Dual-Access Security Architecture):\n").bold = True
    p.add_run("   • 생전/응급 시: 중태·의식불명 발생 시 사전 연명의료 의향서 및 모바일 응급 QR 카드로 의료진 및 대리인 즉시 열람.\n")
    p.add_run("   • 사후 시: 사망진단서 OCR 인증 및 2인 유족 승인(Multi-Sig) 검증 완료 후 256-bit AES 자산 금고 복호화 및 유언 메시지 발송.\n\n")
    
    p.add_run("2. 15개 웹 컴포넌트 원스톱 통합 솔루션 (All-in-One Integration):\n").bold = True
    p.add_run("   • 파편화된 4대 서비스(고이의 장지 비교, 마음부고의 모바일 부고장, iFA의 상속세 계산기, 망고하다의 유언장)를 단 하나의 웹 플랫폼에서 완벽 통합 제공.\n\n")

    p.add_run("3. 고부가가치 B2B2C 세무/상속 컨설팅 수익 모델:\n").bold = True
    p.add_run("   • iFA의 상속 세무 역량과 고이의 가격 투명성을 결합하여, 상속세 시뮬레이션 결과와 연계된 1:1 세무사/변호사 실시간 상담 챗을 통해 높은 고객 생애 가치(LTV) 확보.\n\n")

    p.add_run("4. 프리미엄 이어봄 Design System (Dark Navy & Warm Gold):\n").bold = True
    p.add_run("   • Deep Dark Navy(#1E293B)와 Warm Amber Gold(#D97706) 중심의 신뢰감 높은 1440px 반응형 웹 UX/UI, 실시간 데이터 시각화 차트 및 직관적 뱃지 시스템 적용.")

    # Save DOCX
    doc.save(file_path)
    print("DOCX created successfully at:", file_path)

# --- 2. HTML CONTENT FOR PDF GENERATION ---
html_content = """<!DOCTYPE html>
<html lang="ko">
<head>
<meta charset="UTF-8">
<title>국내 웰다잉 및 장례 테크 4대 서비스 현황 및 이어봄 경쟁력 분석 보고서</title>
<style>
  @import url('https://cdn.jsdelivr.net/gh/orioncactus/pretendard/dist/web/static/pretendard.css');
  
  @page {
    size: A4;
    margin: 14mm 14mm 14mm 14mm;
  }
  
  * {
    box-sizing: border-box;
    -webkit-print-color-adjust: exact !important;
    print-color-adjust: exact !important;
  }
  
  body {
    font-family: 'Pretendard', -apple-system, sans-serif;
    color: #0F172A;
    background-color: #FFFFFF;
    margin: 0;
    padding: 0;
    font-size: 9.2pt;
    line-height: 1.55;
  }

  /* Cover Page */
  .cover {
    min-height: 260mm;
    display: flex;
    flex-direction: column;
    justify-content: space-between;
    padding: 45px 35px;
    background: linear-gradient(135deg, #0F172A 0%, #1E293B 100%);
    color: #FFFFFF;
    border-radius: 12px;
  }
  
  .cover-header {
    border-bottom: 3px solid #D97706;
    padding-bottom: 20px;
  }
  .cover-tag {
    background-color: #D97706;
    color: #FFFFFF;
    padding: 5px 14px;
    font-size: 9.5pt;
    font-weight: 700;
    border-radius: 4px;
    display: inline-block;
    letter-spacing: 1px;
    margin-bottom: 18px;
  }
  .cover-title {
    font-size: 23pt;
    font-weight: 800;
    margin: 0 0 12px 0;
    color: #FFFFFF;
    line-height: 1.35;
  }
  .cover-subtitle {
    font-size: 12pt;
    color: #94A3B8;
    margin: 0;
    font-weight: 400;
  }
  .cover-body {
    margin: 35px 0;
    background: rgba(255, 255, 255, 0.05);
    border: 1px solid rgba(255, 255, 255, 0.1);
    border-radius: 10px;
    padding: 22px;
  }
  .cover-footer {
    border-top: 1px solid #334155;
    padding-top: 20px;
    display: flex;
    justify-content: space-between;
    align-items: flex-end;
    font-size: 9pt;
    color: #CBD5E1;
  }

  .page-break {
    page-break-before: always;
  }

  /* Section Styling */
  .section-title {
    font-size: 13.5pt;
    font-weight: 800;
    color: #1E293B;
    border-left: 5px solid #D97706;
    padding-left: 10px;
    margin-top: 20px;
    margin-bottom: 12px;
  }
  .subsection-title {
    font-size: 10.5pt;
    font-weight: 700;
    color: #0F172A;
    margin-top: 14px;
    margin-bottom: 6px;
  }

  p {
    margin-top: 0;
    margin-bottom: 8px;
    color: #334155;
    text-align: justify;
  }

  /* Key Metrics Grid */
  .metrics-grid {
    display: flex;
    gap: 10px;
    margin: 12px 0;
  }
  .metric-card {
    flex: 1;
    background-color: #F8FAFC;
    border: 1px solid #E2E8F0;
    border-radius: 8px;
    padding: 10px;
    text-align: center;
  }
  .metric-value {
    font-size: 14pt;
    font-weight: 800;
    color: #D97706;
    margin-bottom: 2px;
  }
  .metric-label {
    font-size: 8pt;
    color: #64748B;
    font-weight: 600;
  }

  /* Tables */
  table {
    width: 100%;
    border-collapse: collapse;
    margin: 12px 0;
    font-size: 8.2pt;
  }
  th {
    background-color: #1E293B;
    color: #FFFFFF;
    padding: 7px 8px;
    text-align: left;
    font-weight: 700;
  }
  td {
    border-bottom: 1px solid #E2E8F0;
    padding: 7px 8px;
    color: #334155;
    vertical-align: top;
  }
  tr:nth-child(even) td {
    background-color: #F8FAFC;
  }

  /* Badge Tokens */
  .badge {
    display: inline-block;
    padding: 2px 5px;
    border-radius: 4px;
    font-size: 7.2pt;
    font-weight: 700;
  }
  .badge-primary { background: #E0F2FE; color: #0369A1; }
  .badge-success { background: #ECFDF5; color: #047857; }
  .badge-warning { background: #FEF3C7; color: #B45309; }
  .badge-danger { background: #FEE2E2; color: #B91C1C; }

  /* Callout Box */
  .callout {
    background-color: #FFFBEB;
    border-left: 4px solid #D97706;
    padding: 12px 15px;
    margin: 14px 0;
    border-radius: 0 8px 8px 0;
    font-size: 8.8pt;
  }
  .callout-title {
    font-weight: 700;
    color: #B45309;
    font-size: 9.5pt;
    margin-bottom: 4px;
  }
  
  ul, ol {
    margin-top: 4px;
    margin-bottom: 8px;
    padding-left: 18px;
  }
  li {
    margin-bottom: 3px;
  }
</style>
</head>
<body>

<!-- COVER PAGE -->
<div class="cover">
  <div class="cover-header">
    <div class="cover-tag">전략 기획 및 경쟁력 분석 보고서</div>
    <div class="cover-title">국내 웰다잉·장례 테크 4대 서비스 현황 및 이어봄 플랫폼 경쟁력 분석 보고서</div>
    <div class="cover-subtitle">실제 4대 웹사이트(고이, 마음부고, iFA 엔딩플랜, 망고하다) 딥다이브 분석 및 이어봄 차별화 가치 검증</div>
  </div>
  
  <div class="cover-body">
    <p style="color: #F8FAFC; font-size: 10pt; line-height: 1.8; margin: 0;">
      본 보고서는 대한민국 초고령사회 진입에 따른 웰다잉(Well-Dying) 및 디지털 엔딩(Digital Ending) 시장의 생태계를 정밀 조사하고,<br>
      대표 장례 플랫폼 <strong>'고이(goifuneral.co.kr)'</strong>, 무료 모바일 부고장 <strong>'마음부고(mbugo.com)'</strong>, 인슈테크 GA <strong>'iFA 엔딩플랜(ifa.co.kr)'</strong>, 웰다잉 전문 스타트업 <strong>'망고하다(mangohada.com)'</strong> 4개 홈페이지를 직접 분석하여<br>
      기존 서비스들의 강점과 한계를 파악하고 <strong>이어봄 토탈 케어 플랫폼의 독보적인 4대 경쟁 우위</strong>를 제시합니다.
    </p>
  </div>

  <div class="cover-footer">
    <div>
      <div><strong>발행 기관:</strong> 이어봄 토탈 케어 서비스 전략 기획팀</div>
      <div><strong>발행 일자:</strong> 2026년 07월 30일</div>
    </div>
    <div style="text-align: right;">
      <div><strong>문서 분류:</strong> 대외비 (Confidential Document)</div>
      <div><strong>플랫폼:</strong> 이어봄 & LastLetter Care</div>
    </div>
  </div>
</div>

<div class="page-break"></div>

<!-- SECTION 1 -->
<div class="section-title">1. 국내 시장 환경 및 배경 분석 (Market Overview)</div>

<p>
  대한민국은 2025년을 기점으로 65세 이상 고령 인구 비율이 20%를 돌파하는 <strong>'초고령사회'</strong>에 공식 진입하였습니다. 
  이에 따라 사후 케어와 유산 정리를 당사자가 존엄하게 준비하는 <strong>웰다잉(Well-Dying)</strong> 및 <strong>디지털 엔딩(Digital Ending)</strong> 서비스가 
  상조 산업을 넘어 IT 기반의 라이프 테크 분야로 성장하고 있습니다.
</p>

<div class="metrics-grid">
  <div class="metric-card">
    <div class="metric-value">20%+</div>
    <div class="metric-label">초고령사회 고령 인구 비율</div>
  </div>
  <div class="metric-card">
    <div class="metric-value">120억 원</div>
    <div class="metric-label">고이장례연구소 누적 투자액</div>
  </div>
  <div class="metric-card">
    <div class="metric-value">67억 원</div>
    <div class="metric-label">후불제 상조 스타트업 연 매출</div>
  </div>
  <div class="metric-card">
    <div class="metric-value">2,200명+</div>
    <div class="metric-label">iFA 인슈테크 GA 전역 조직</div>
  </div>
</div>

<div class="subsection-title">1.1 기존 시장의 3대 핵심 페인 포인트</div>
<ul>
  <li><strong>선불 상조 예치금 리스크:</strong> 기존 상조업체의 선불 납입금 유동성 위험 및 폐업 불안감 지속.</li>
  <li><strong>장례 비용 비대칭성:</strong> 장례식장, 용품, 봉안당/수목장 장지 비용의 사전 비교 시스템 부족.</li>
  <li><strong>서비스 파편화:</strong> 엔딩노트, 부고장, 연명의료 의향서, 상속세 계산, SNS 계정 정리가 분산되어 통합 케어 서비스 공백 존재.</li>
</ul>

<!-- SECTION 2 -->
<div class="section-title">2. 국내 4대 핵심 서비스 정밀 분석 (Live Website Deep Dive)</div>

<table>
  <thead>
    <tr>
      <th style="width: 14%;">구분</th>
      <th style="width: 20%;">웹사이트 및 서비스명</th>
      <th style="width: 42%;">주요 서비스 특징 및 핵심 모델</th>
      <th style="width: 24%;">시장 내 평가 및 한계점</th>
    </tr>
  </thead>
  <tbody>
    <tr>
      <td><strong>장례/장지 테크</strong></td>
      <td><strong>고이 (Goi)</strong><br><a href="https://www.goifuneral.co.kr/">goifuneral.co.kr</a><br><span class="badge badge-success">누적 120억 투자</span></td>
      <td>• <strong>'100원 상조'</strong> 후불제 모델 (월 100원 계약금, 서비스 발생 후 정산)<br>• 미사용 물품 100% 반품/환불 정찰제<br>• 전국 장례식장/봉안당 맞춤 실시간 견적 및 24h 컨시어지<br>• 2025년 매출 67억 원 달성 (설립 1년 대비 26배 성장)</td>
      <td>사후 의전 및 장지 시설 견적 비교에 특화됨. 생전 응급 QR 카드, 암호화 엔딩노트, 상속세 계산기 기능 미제공.</td>
    </tr>
    <tr>
      <td><strong>모바일 부고장</strong></td>
      <td><strong>마음부고 (mbugo)</strong><br><a href="https://www.mbugo.com/">mbugo.com</a><br><span class="badge badge-primary">100% 무료 서비스</span></td>
      <td>• 회원가입 및 앱 설치 없는 1분 모바일 부고장/감사장 무료 제작<br>• 카카오톡 및 SMS 부고 전달, 계좌번호 및 오시는 길 안내<br>• 전국 장례식장 직통 전화 연결 및 지자체 부고 검색</td>
      <td>단순 모바일 부고장 생성/전달 도구로, 웰다잉 엔딩노트, 장지 선택, 상속 세무, 사후 유품 정리 기능 없음.</td>
    </tr>
    <tr>
      <td><strong>상속/인슈테크</strong></td>
      <td><strong>iFA 엔딩플랜</strong><br><a href="https://www.ifa.co.kr/">ifa.co.kr</a><br><span class="badge badge-warning">인슈테크 GA</span></td>
      <td>• 상속·증여 전문 연구소 운영 (분석 ➔ 해법 ➔ 제안 ➔ 공증 4단계)<br>• 서울 15.5억 아파트 기준 1.25억 상속세 모의 계산기 (Tax_Main.aspx)<br>• 사망보험금 레버리지를 활용한 절세 및 자산 이전 컨설팅 (특허 4건)</td>
      <td>보험 상품 판매 및 오프라인 상속 컨설팅 마케팅 수단 성격. 디지털 유품 정리, 응급 QR 카드, 장지 실시간 비교 부재.</td>
    </tr>
    <tr>
      <td><strong>웰다잉 스타트업</strong></td>
      <td><strong>망고하다</strong><br><a href="https://www.mangohada.com/">mangohada.com</a><br><span class="badge badge-success">웰다잉 스타트업</span></td>
      <td>• 정서적 유언 기록 및 버킷리스트 작성 웰다잉 전문 플랫폼<br>• 장례 의전 전문 기업(엔딩스케치) MOU를 통해 온라인 기록의 실물 장례 연결</td>
      <td>스토리텔링 중심 유언 기록 플랫폼으로, 256-bit AES 기술적 보안 금고, 2인 유족 승인(Multi-Sig), 상속세 시뮬레이터 미흡.</td>
    </tr>
  </tbody>
</table>

<div class="page-break"></div>

<!-- SECTION 3 -->
<div class="section-title">3. 이어봄 플랫폼 핵심 경쟁력 비교 (Feature Comparison Matrix)</div>

<table>
  <thead>
    <tr>
      <th style="width: 25%;">비교 핵심 기능</th>
      <th style="width: 20%;">이어봄 (본 서비스)</th>
      <th style="width: 15%;">고이 (Goi)</th>
      <th style="width: 13%;">마음부고</th>
      <th style="width: 13%;">iFA 엔딩플랜</th>
      <th style="width: 14%;">망고하다</th>
    </tr>
  </thead>
  <tbody>
    <tr>
      <td><strong>서비스 통합 범위</strong></td>
      <td><span class="badge badge-success">생전~사후 올인원</span></td>
      <td>사후 의전/장지</td>
      <td>모바일 부고장</td>
      <td>상속세/보험</td>
      <td>유언/의전 연계</td>
    </tr>
    <tr>
      <td><strong>생전/응급 의료 QR 카드</strong></td>
      <td><span class="badge badge-success">지원 (사전연명의료 연동)</span></td>
      <td><span class="badge badge-danger">미지원</span></td>
      <td><span class="badge badge-danger">미지원</span></td>
      <td><span class="badge badge-danger">미지원</span></td>
      <td><span class="badge badge-danger">미지원</span></td>
    </tr>
    <tr>
      <td><strong>256-bit AES & Multi-Sig 금고</strong></td>
      <td><span class="badge badge-success">지원 (유족 2인 검증)</span></td>
      <td><span class="badge badge-danger">미지원</span></td>
      <td><span class="badge badge-danger">미지원</span></td>
      <td><span class="badge badge-danger">미지원</span></td>
      <td><span class="badge badge-danger">미지원</span></td>
    </tr>
    <tr>
      <td><strong>전국 8대 권역 장지 비교</strong></td>
      <td><span class="badge badge-success">지원 (맞춤 필터 & 예약)</span></td>
      <td><span class="badge badge-primary">지원 (실시간 견적)</span></td>
      <td>미지원</td>
      <td>미지원</td>
      <td>미지원</td>
    </tr>
    <tr>
      <td><strong>상속세 자동 시뮬레이터</strong></td>
      <td><span class="badge badge-success">지원 (자산/부채 차트)</span></td>
      <td><span class="badge badge-danger">미지원</span></td>
      <td>미지원</td>
      <td><span class="badge badge-primary">지원 (모의계산기)</span></td>
      <td>미지원</td>
    </tr>
    <tr>
      <td><strong>디지털 유품 (SNS 계정) 정리</strong></td>
      <td><span class="badge badge-success">지원 (3/8건 진행 트래킹)</span></td>
      <td><span class="badge badge-danger">미지원</span></td>
      <td>미지원</td>
      <td>미지원</td>
      <td>미지원</td>
    </tr>
    <tr>
      <td><strong>1:1 전문가 상담 챗</strong></td>
      <td><span class="badge badge-success">지원 (실시간 예약)</span></td>
      <td>24h 콜센터</td>
      <td>미지원</td>
      <td>오프라인 상담</td>
      <td>미지원</td>
    </tr>
  </tbody>
</table>

<div class="callout">
  <div class="callout-title">💡 이어봄 플랫폼의 4대 독보적 핵심 우위 (Core Competitive Edges)</div>
  <ol style="margin: 5px 0 0 15px; padding: 0;">
    <li><strong>생전-응급-사후 3-Tier 이원화 보안 체계:</strong> 생전 중태/응급 시 사전 연명의료 의향서 QR 카드로 즉시 열람 지원 + 사후 사망진단서 OCR 및 유족 2인 승인(Multi-Sig) 후 256-bit AES 자산 금고 복호화 공개.</li>
    <li><strong>15개 웹 컴포넌트 원스톱 통합 솔루션 (All-in-One):</strong> 파편화되어 있는 엔딩노트, 장지 비교, 상속세 계산, 디지털 유품 정리, 웰다잉 교육, 24시 긴급가이드를 단 하나의 플랫폼에서 완벽 처리.</li>
    <li><strong>고부가가치 B2B2C 세무/상속 컨설팅 모델:</strong> iFA의 세무 역량과 고이의 가격 투명성을 결합하여, 상속세 시뮬레이션 결과와 연계된 1:1 세무사/변호사 실시간 상담 챗을 통해 높은 LTV 및 고수수료 매출 확보.</li>
    <li><strong>프리미엄 Web UX/UI 디자인 시스템 (이어봄 Design System):</strong> Dark Navy (`#1E293B`) & Warm Gold (`#D97706`) 중심의 신뢰감 높은 1440px 반응형 그리드 및 데이터 시각화 차트 적용.</li>
  </ol>
</div>

<!-- SECTION 4 -->
<div class="section-title">4. 결론 및 시장 침투 전략 (Strategic Roadmap)</div>

<div class="subsection-title">4.1 단계별 실행 로드맵</div>
<ol>
  <li><strong>1단계 (고이 모델 수용 & 장지 파트너십 구축):</strong> 고이(Goi)의 후불제 가격 정찰제 장점을 흡수하고 전국 8대 권역 봉안당·수목장 맞춤 비교 서비스 구축.</li>
  <li><strong>2단계 (응급 QR & 엔딩노트 보급 유저 락인):</strong> 시니어 및 3040 자녀 세대를 대상으로 사전 연명의료 의향서 응급 QR 카드 및 256-bit 엔딩노트 무료 보급.</li>
  <li><strong>3단계 (상속 세무 B2B 파트너십 & 구독 수익화):</strong> 세무법인·법무법인 연계 상속 케어 수수료 및 디지털 유품 일괄 해지/정리 대행 서비스 프리미엄 플랜 상용화.</li>
</ol>

<div style="margin-top: 35px; border-top: 2px solid #E2E8F0; padding-top: 15px; text-align: center; color: #94A3B8; font-size: 8.5pt;">
  이어봄 Total Care Platform — Confidential Research Report ⓒ 2026 이어봄 All Rights Reserved.
</div>

</body>
</html>
"""

# Output paths exclusively inside REPORTS_DIR
html_path = os.path.join(REPORTS_DIR, "이어봄_경쟁_서비스_분석_보고서.html")
docx_path = os.path.join(REPORTS_DIR, "이어봄_경쟁_서비스_분석_보고서.docx")
pdf_path = os.path.join(REPORTS_DIR, "이어봄_경쟁_서비스_분석_보고서.pdf")

# 1. Create DOCX Report
create_docx_report(docx_path)

# 2. Save HTML & Convert to PDF via Edge
with open(html_path, "w", encoding="utf-8") as f:
    f.write(html_content)

edge_path = r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe"
cmd = [
    edge_path,
    "--headless",
    "--disable-gpu",
    "--no-pdf-header-footer",
    f"--print-to-pdf={pdf_path}",
    html_path
]

res = subprocess.run(cmd, capture_output=True, text=True)

print("\n--- SINGLE REPORT DIRECTORY OUTPUT ---")
print("Target Directory:", REPORTS_DIR)
print("DOCX File:", os.path.exists(docx_path), os.path.getsize(docx_path) if os.path.exists(docx_path) else 0, "bytes")
print("PDF File:", os.path.exists(pdf_path), os.path.getsize(pdf_path) if os.path.exists(pdf_path) else 0, "bytes")
